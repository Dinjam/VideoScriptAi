# VideoScript AI — Faster-Whisper Edition

"""
VideoScript AI
Video → Transcript → Learning Script

Single-file Streamlit application
Optimized for Streamlit Community Cloud / CPU.

Install locally:
    pip install -r requirements.txt

Run:
    streamlit run video_converter.py

Requires:
    FFmpeg

Optional:
    GROQ_API_KEY for AI Learning Script and Summary.
"""

import os
import re
import io
import shutil
import tempfile
import subprocess
from pathlib import Path

import streamlit as st


# ============================================================
# OPTIONAL / REQUIRED DEPENDENCIES
# ============================================================

try:
    from faster_whisper import WhisperModel
except Exception:
    WhisperModel = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
    )
    from reportlab.lib.styles import getSampleStyleSheet
except Exception:
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    getSampleStyleSheet = None

try:
    from groq import Groq
except Exception:
    Groq = None


# ============================================================
# CONFIG
# ============================================================

APP_NAME = "VideoScript AI"

MAX_UPLOAD_MB = 2048

# Faster-Whisper model:
# tiny   = fastest
# base   = recommended for cloud CPU
# small  = better accuracy but heavier
# medium = not recommended on free Streamlit Cloud
DEFAULT_MODEL = "base"

# CPU configuration
DEVICE = "cpu"
COMPUTE_TYPE = "int8"

st.set_page_config(
    page_title=APP_NAME,
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ============================================================
# STYLE
# ============================================================

st.markdown(
    """
    <style>
        .stApp {
            background: #0b1020;
        }

        [data-testid="stHeader"] {
            background: rgba(0,0,0,0);
        }

        .hero {
            padding: 28px 30px;
            border: 1px solid rgba(255,255,255,.08);
            border-radius: 22px;
            background: linear-gradient(135deg,#141c35,#101629);
            margin-bottom: 20px;
        }

        .hero h1 {
            margin: 0;
            font-size: 42px;
        }

        .hero p {
            color: #aeb8d0;
            font-size: 17px;
        }

        .card {
            padding: 18px;
            border-radius: 18px;
            border: 1px solid rgba(255,255,255,.08);
            background: #11182b;
            margin-bottom: 14px;
        }

        .metric {
            padding: 15px;
            border-radius: 16px;
            background: #151e35;
            text-align: center;
        }

        .muted {
            color: #9ca8c4;
        }

        .success {
            padding: 12px 16px;
            border-radius: 12px;
            background: rgba(34,197,94,.12);
            border: 1px solid rgba(34,197,94,.3);
        }
    </style>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# HELPERS
# ============================================================

def ffmpeg_available():
    return shutil.which("ffmpeg") is not None


def get_ffmpeg_path():
    return shutil.which("ffmpeg")


def format_seconds(seconds):
    seconds = max(0, int(seconds))

    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60

    return f"{h:02d}:{m:02d}:{s:02d}"


def clean_text(text):
    """
    Lightweight transcript cleanup.
    Does not generate new information.
    """

    if not text:
        return ""

    text = re.sub(r"\s+", " ", text).strip()

    fillers = [
        r"\b(eh+|eee+|emm+|hmm+|um+|uh+)\b",
        r"\b(ya+ ya+ ya+)\b",
    ]

    for pattern in fillers:
        text = re.sub(pattern, "", text, flags=re.I)

    text = re.sub(r"\s+([,.!?;:])", r"\1", text)

    text = re.sub(
        r"([,.!?])([A-Za-zÀ-ÿ])",
        r"\1 \2",
        text,
    )

    text = re.sub(r"\s{2,}", " ", text)

    return text.strip()


# ============================================================
# SUBTITLE EXPORT
# ============================================================

def segments_to_srt(segments):

    def timestamp(sec):
        sec = max(0, float(sec))

        hours = int(sec // 3600)
        minutes = int((sec % 3600) // 60)
        seconds = int(sec % 60)
        milliseconds = int(
            round((sec - int(sec)) * 1000)
        )

        if milliseconds >= 1000:
            seconds += 1
            milliseconds = 0

        return (
            f"{hours:02d}:"
            f"{minutes:02d}:"
            f"{seconds:02d},"
            f"{milliseconds:03d}"
        )

    blocks = []

    for i, seg in enumerate(segments, 1):

        blocks.append(
            f"{i}\n"
            f"{timestamp(seg['start'])} --> "
            f"{timestamp(seg['end'])}\n"
            f"{clean_text(seg['text'])}\n"
        )

    return "\n".join(blocks)


def segments_to_vtt(segments):

    def timestamp(sec):
        sec = max(0, float(sec))

        hours = int(sec // 3600)
        minutes = int((sec % 3600) // 60)
        seconds = int(sec % 60)
        milliseconds = int(
            round((sec - int(sec)) * 1000)
        )

        return (
            f"{hours:02d}:"
            f"{minutes:02d}:"
            f"{seconds:02d}."
            f"{milliseconds:03d}"
        )

    lines = [
        "WEBVTT",
        "",
    ]

    for seg in segments:

        lines += [
            (
                f"{timestamp(seg['start'])} --> "
                f"{timestamp(seg['end'])}"
            ),
            clean_text(seg["text"]),
            "",
        ]

    return "\n".join(lines)


def transcript_markdown(title, segments, clean=False):

    lines = [
        f"# {title}",
        "",
        "## Transkrip",
        "",
    ]

    for seg in segments:

        text = (
            clean_text(seg["text"])
            if clean
            else seg["text"].strip()
        )

        lines.append(
            f"**[{format_seconds(seg['start'])}]** {text}"
        )

        lines.append("")

    return "\n".join(lines)


# ============================================================
# DOCUMENT EXPORT
# ============================================================

def make_docx(title, text):

    if Document is None:
        raise RuntimeError(
            "python-docx belum terpasang."
        )

    doc = Document()

    doc.add_heading(title, 0)

    for block in text.split("\n\n"):

        if block.startswith("# "):

            doc.add_heading(
                block[2:],
                level=1,
            )

        elif block.startswith("## "):

            doc.add_heading(
                block[3:],
                level=2,
            )

        else:

            doc.add_paragraph(block)

    output = io.BytesIO()

    doc.save(output)

    return output.getvalue()


def make_pdf(title, text):

    if SimpleDocTemplate is None:
        raise RuntimeError(
            "reportlab belum terpasang."
        )

    output = io.BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
    )

    styles = getSampleStyleSheet()

    story = [
        Paragraph(
            title,
            styles["Title"],
        ),
        Spacer(1, 12),
    ]

    for block in text.split("\n\n"):

        safe = (
            block
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

        if safe.startswith("# "):

            story.append(
                Paragraph(
                    safe[2:],
                    styles["Heading1"],
                )
            )

        elif safe.startswith("## "):

            story.append(
                Paragraph(
                    safe[3:],
                    styles["Heading2"],
                )
            )

        else:

            story.append(
                Paragraph(
                    safe.replace(
                        "\n",
                        "<br/>",
                    ),
                    styles["BodyText"],
                )
            )

        story.append(
            Spacer(1, 8)
        )

    doc.build(story)

    return output.getvalue()


# ============================================================
# FFMPEG
# ============================================================

def extract_audio(video_path, audio_path):

    ffmpeg_path = get_ffmpeg_path()

    if not ffmpeg_path:
        raise RuntimeError(
            "FFmpeg tidak ditemukan."
        )

    command = [
        ffmpeg_path,
        "-y",
        "-i",
        str(video_path),
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        "-c:a",
        "pcm_s16le",
        str(audio_path),
    ]

    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
    )

    if result.returncode != 0:

        raise RuntimeError(
            "FFmpeg gagal memproses file:\n\n"
            + result.stderr[-3000:]
        )


# ============================================================
# FASTER-WHISPER MODEL
# ============================================================

@st.cache_resource(
    show_spinner=False
)
def load_whisper_model(model_name):

    if WhisperModel is None:

        raise RuntimeError(
            "Faster-Whisper belum terpasang.\n\n"
            "Install dengan:\n"
            "pip install faster-whisper"
        )

    return WhisperModel(
        model_name,
        device=DEVICE,
        compute_type=COMPUTE_TYPE,
        cpu_threads=4,
        num_workers=1,
    )


# ============================================================
# TRANSCRIPTION
# ============================================================

def transcribe(
    audio_path,
    model_name,
    language,
):

    model = load_whisper_model(
        model_name
    )

    whisper_language = None

    if language != "auto":
        whisper_language = language

    segments_generator, info = model.transcribe(
        str(audio_path),
        language=whisper_language,
        task="transcribe",
        beam_size=1,
        best_of=1,
        temperature=0,
        vad_filter=True,
        condition_on_previous_text=False,
    )

    segments = []

    total_duration = float(
        getattr(
            info,
            "duration",
            0,
        )
        or 0
    )

    progress_bar = st.progress(0)

    status = st.empty()

    for segment in segments_generator:

        text = segment.text.strip()

        segments.append(
            {
                "start": float(
                    segment.start
                ),
                "end": float(
                    segment.end
                ),
                "text": text,
            }
        )

        if total_duration > 0:

            progress = min(
                int(
                    (
                        segment.end
                        / total_duration
                    )
                    * 100
                ),
                100,
            )

            progress_bar.progress(
                progress
            )

            status.info(
                "🎙️ Menganalisis audio... "
                f"{progress}%"
            )

    progress_bar.progress(100)

    status.success(
        "✅ Transkripsi selesai!"
    )

    raw_text = " ".join(
        seg["text"]
        for seg in segments
    ).strip()

    detected_language = getattr(
        info,
        "language",
        "unknown",
    )

    return (
        raw_text,
        segments,
        detected_language,
    )


# ============================================================
# GROQ AI
# ============================================================

def generate_ai_learning_script(
    transcript,
    mode="learning",
):

    api_key = os.getenv(
        "GROQ_API_KEY"
    )

    if not api_key or Groq is None:
        return None

    client = Groq(
        api_key=api_key
    )

    if mode == "summary":

        instruction = """
Buat ringkasan pembelajaran dari transkrip berikut.

Gunakan hanya informasi yang ada di transkrip.

Format:

# Ringkasan

## Gambaran Umum

## Poin Penting

## Kesimpulan
"""

    else:

        instruction = """
Ubah transkrip berikut menjadi learning script
yang rapi dan mudah dipelajari.

Jangan mengarang informasi baru.

Pertahankan istilah penting dan makna asli.

Format:

# Judul Materi

## Gambaran Umum

## Konsep Utama

## Contoh

## Hal yang Harus Diingat

## Kesimpulan

Jika suatu bagian tidak tersedia dalam transkrip,
jangan mengada-adakan.
"""

    prompt = (
        instruction
        + "\n\nTRANSKRIP:\n"
        + transcript
    )

    response = client.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[
            {
                "role": "system",
                "content": (
                    "Kamu adalah editor materi "
                    "pembelajaran. Tulis dalam "
                    "Bahasa Indonesia yang jelas, "
                    "ringkas, dan akurat."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        temperature=0.2,
    )

    return (
        response
        .choices[0]
        .message
        .content
    )


# ============================================================
# SESSION STATE
# ============================================================

defaults = {

    "segments": [],

    "raw_transcript": "",

    "clean_transcript": "",

    "learning_script": "",

    "summary": "",

    "source_name": "",

    "detected_language": "",

    "processed": False,
}

for key, value in defaults.items():

    st.session_state.setdefault(
        key,
        value,
    )


# ============================================================
# HEADER
# ============================================================

st.markdown(
    """
    <div class="hero">

        <h1>🎬 VideoScript AI</h1>

        <p>
            Ubah video/audio menjadi transkrip,
            catatan belajar, ringkasan, subtitle,
            dan dokumen yang mudah dibaca.
        </p>

    </div>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# SIDEBAR
# ============================================================

with st.sidebar:

    st.header(
        "⚙️ Pengaturan"
    )

    model_name = st.selectbox(
        "Whisper model",
        [
            "tiny",
            "base",
            "small",
        ],
        index=[
            "tiny",
            "base",
            "small",
        ].index(DEFAULT_MODEL),
        help=(
            "Base direkomendasikan untuk "
            "Streamlit Cloud."
        ),
    )

    language = st.selectbox(
        "Bahasa",
        [
            (
                "auto",
                "Deteksi otomatis",
            ),
            (
                "id",
                "Indonesia",
            ),
            (
                "en",
                "English",
            ),
            (
                "ja",
                "Japanese",
            ),
        ],
        format_func=lambda x: x[1],
    )[0]

    clean_enabled = st.checkbox(
        "Bersihkan transkrip",
        True,
    )

    ai_enabled = st.checkbox(
        "Buat Learning Script dengan AI",
        bool(
            os.getenv(
                "GROQ_API_KEY"
            )
        ),
    )

    st.divider()

    st.caption(
        "Faster-Whisper Edition"
    )

    st.caption(
        "CPU: int8"
    )

    st.caption(
        "Video → Audio → Whisper → AI → Export"
    )


# ============================================================
# SYSTEM STATUS
# ============================================================

if not ffmpeg_available():

    st.error(
        "FFmpeg belum ditemukan di server."
    )

else:

    st.success(
        "FFmpeg siap digunakan."
    )


st.info(
    f"🧠 Model: **{model_name}**  |  "
    f"⚙️ Device: **{DEVICE.upper()}**  |  "
    f"Precision: **{COMPUTE_TYPE}**"
)


# ============================================================
# UPLOAD
# ============================================================

uploaded = st.file_uploader(
    "Upload video/audio",
    type=[
        "mp4",
        "mkv",
        "mov",
        "avi",
        "webm",
        "mp3",
        "wav",
        "m4a",
        "aac",
        "flac",
        "ogg",
    ],
    help=(
        f"Ukuran maksimum konfigurasi aplikasi: "
        f"{MAX_UPLOAD_MB} MB."
    ),
)


if uploaded:

    size_mb = (
        uploaded.size
        / (1024 * 1024)
    )

    st.markdown(
        f"""
        <div class="card">

            <b>📁 {uploaded.name}</b>
            <br>

            <span class="muted">
                {size_mb:.1f} MB
            </span>

        </div>
        """,
        unsafe_allow_html=True,
    )

    if size_mb > MAX_UPLOAD_MB:

        st.error(
            "File terlalu besar."
        )

    else:

        col1, col2 = st.columns(
            [1, 1]
        )

        with col1:

            process = st.button(
                "🚀 Mulai Transkripsi",
                type="primary",
                use_container_width=True,
            )

        with col2:

            reset = st.button(
                "↻ Reset",
                use_container_width=True,
            )

        if reset:

            for key, value in defaults.items():

                st.session_state[key] = value

            st.rerun()

        if process:

            with tempfile.TemporaryDirectory() as tmp:

                source = (
                    Path(tmp)
                    / uploaded.name
                )

                source.write_bytes(
                    uploaded.getbuffer()
                )

                audio = (
                    Path(tmp)
                    / "audio.wav"
                )

                try:

                    suffix = (
                        source.suffix.lower()
                    )

                    audio_extensions = {
                        ".mp3",
                        ".wav",
                        ".m4a",
                        ".aac",
                        ".flac",
                        ".ogg",
                    }

                    if suffix in audio_extensions:

                        audio = source

                        st.info(
                            "🎵 File audio langsung digunakan."
                        )

                    else:

                        st.info(
                            "🎵 Mengekstrak audio..."
                        )

                        extract_audio(
                            source,
                            audio,
                        )

                    st.success(
                        "✅ Audio siap."
                    )

                    st.info(
                        "🧠 Memuat Faster-Whisper "
                        f"model **{model_name}**..."
                    )

                    raw, segments, detected = (
                        transcribe(
                            audio,
                            model_name,
                            language,
                        )
                    )

                    clean = " ".join(
                        clean_text(
                            seg["text"]
                        )
                        for seg in segments
                    ).strip()

                    st.session_state.segments = (
                        segments
                    )

                    st.session_state.raw_transcript = (
                        raw.strip()
                    )

                    st.session_state.clean_transcript = (
                        clean
                        if clean_enabled
                        else raw.strip()
                    )

                    st.session_state.source_name = (
                        uploaded.name
                    )

                    st.session_state.detected_language = (
                        detected
                    )

                    st.session_state.processed = (
                        True
                    )

                    if ai_enabled:

                        st.info(
                            "✨ Membuat Learning Script..."
                        )

                        learning = (
                            generate_ai_learning_script(
                                st.session_state.clean_transcript,
                                "learning",
                            )
                        )

                        st.session_state.learning_script = (
                            learning or ""
                        )

                    else:

                        st.session_state.learning_script = (
                            ""
                        )

                    st.success(
                        "🎉 Semua proses selesai!"
                    )

                    st.rerun()

                except Exception as e:

                    st.error(
                        f"❌ Proses gagal: {e}"
                    )


# ============================================================
# RESULTS
# ============================================================

if st.session_state.processed:

    st.divider()

    segments = (
        st.session_state.segments
    )

    total_duration = (
        segments[-1]["end"]
        if segments
        else 0
    )

    words = len(
        st.session_state
        .clean_transcript
        .split()
    )

    c1, c2, c3, c4 = st.columns(4)

    with c1:

        st.markdown(
            f"""
            <div class="metric">
                <b>Durasi</b><br>
                {format_seconds(total_duration)}
            </div>
            """,
            unsafe_allow_html=True,
        )

    with c2:

        st.markdown(
            f"""
            <div class="metric">
                <b>Segmen</b><br>
                {len(segments)}
            </div>
            """,
            unsafe_allow_html=True,
        )

    with c3:

        st.markdown(
            f"""
            <div class="metric">
                <b>Kata</b><br>
                {words:,}
            </div>
            """,
            unsafe_allow_html=True,
        )

    with c4:

        st.markdown(
            f"""
            <div class="metric">
                <b>Bahasa</b><br>
                {st.session_state
                    .detected_language
                    .upper()}
            </div>
            """,
            unsafe_allow_html=True,
        )

    tabs = st.tabs(
        [
            "📝 Transkrip",
            "📚 Learning Script",
            "⚡ Ringkasan",
            "⏱️ Timestamp",
            "📤 Export",
        ]
    )

    # ========================================================
    # TRANSCRIPT
    # ========================================================

    with tabs[0]:

        st.subheader(
            "Transkrip Bersih"
        )

        edited = st.text_area(
            "Edit transkrip",
            st.session_state.clean_transcript,
            height=500,
        )

        st.session_state.clean_transcript = (
            edited
        )

    # ========================================================
    # LEARNING SCRIPT
    # ========================================================

    with tabs[1]:

        st.subheader(
            "Learning Script"
        )

        if st.session_state.learning_script:

            st.markdown(
                st.session_state.learning_script
            )

        else:

            st.info(
                "Learning Script AI belum dibuat. "
                "Aktifkan opsi AI dan proses video "
                "kembali."
            )

    # ========================================================
    # SUMMARY
    # ========================================================

    with tabs[2]:

        st.subheader(
            "Ringkasan"
        )

        if st.button(
            "✨ Buat Ringkasan AI"
        ):

            with st.spinner(
                "Membuat ringkasan..."
            ):

                summary = (
                    generate_ai_learning_script(
                        st.session_state.clean_transcript,
                        "summary",
                    )
                )

            if summary:

                st.session_state.summary = (
                    summary
                )

                st.rerun()

            else:

                st.warning(
                    "GROQ_API_KEY atau library "
                    "groq belum tersedia."
                )

        if st.session_state.summary:

            st.markdown(
                st.session_state.summary
            )

    # ========================================================
    # TIMESTAMP
    # ========================================================

    with tabs[3]:

        st.subheader(
            "Transkrip dengan Timestamp"
        )

        for seg in segments:

            text = clean_text(
                seg["text"]
            )

            st.markdown(
                f"**[{format_seconds(seg['start'])}]** "
                f"{text}"
            )

    # ========================================================
    # EXPORT
    # ========================================================

    with tabs[4]:

        st.subheader(
            "Download"
        )

        title = Path(
            st.session_state.source_name
        ).stem

        md = transcript_markdown(
            title,
            segments,
            clean=True,
        )

        raw_txt = (
            st.session_state.raw_transcript
        )

        clean_txt = (
            st.session_state.clean_transcript
        )

        srt = segments_to_srt(
            segments
        )

        vtt = segments_to_vtt(
            segments
        )

        col1, col2, col3 = st.columns(3)

        with col1:

            st.download_button(
                "⬇️ TXT",
                clean_txt,
                file_name=(
                    f"{title}_transcript.txt"
                ),
                mime="text/plain",
                use_container_width=True,
            )

        with col2:

            st.download_button(
                "⬇️ Markdown",
                md,
                file_name=(
                    f"{title}_transcript.md"
                ),
                mime="text/markdown",
                use_container_width=True,
            )

        with col3:

            st.download_button(
                "⬇️ SRT",
                srt,
                file_name=f"{title}.srt",
                mime="application/x-subrip",
                use_container_width=True,
            )

        col4, col5, col6 = st.columns(3)

        with col4:

            st.download_button(
                "⬇️ VTT",
                vtt,
                file_name=f"{title}.vtt",
                mime="text/vtt",
                use_container_width=True,
            )

        with col5:

            try:

                docx_data = make_docx(
                    title,
                    md,
                )

                st.download_button(
                    "⬇️ DOCX",
                    docx_data,
                    file_name=(
                        f"{title}.docx"
                    ),
                    mime=(
                        "application/"
                        "vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                    use_container_width=True,
                )

            except Exception:

                st.button(
                    "DOCX — install python-docx",
                    disabled=True,
                    use_container_width=True,
                )

        with col6:

            try:

                pdf_data = make_pdf(
                    title,
                    md,
                )

                st.download_button(
                    "⬇️ PDF",
                    pdf_data,
                    file_name=f"{title}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )

            except Exception:

                st.button(
                    "PDF — install reportlab",
                    disabled=True,
                    use_container_width=True,
                )

        st.divider()

        st.subheader(
            "Learning Script Export"
        )

        if st.session_state.learning_script:

            st.download_button(
                "📚 Download Learning Script",
                st.session_state.learning_script,
                file_name=(
                    f"{title}_learning_script.md"
                ),
                mime="text/markdown",
                use_container_width=True,
            )


# ============================================================
# FOOTER
# ============================================================

st.divider()

st.caption(
    "VideoScript AI — Faster-Whisper Edition. "
    "Gunakan hanya konten yang kamu punya hak "
    "untuk transkripsikan."
)

